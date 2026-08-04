"""Erzeugt ein Word-Dokument mit allen Venues/Personen/Ensembles ohne Bild
(photo_url leer) — reine Referenz-Übersicht für die Redaktion, schreibt
nichts in die Datenbank. Für den tatsächlichen Bild-Beschaffungsprozess
siehe backend/supabase/functions/enrich-entity-images/.
"""

import argparse
import json
import logging
import os
import sys
import urllib.parse
import urllib.request
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger("build_missing_images_doc")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT = "E8EEF5"
WHITE = "FFFFFF"
PAGE_SIZE = 1000


def load_env(path: Path) -> dict:
    values = {}
    if path.exists():
        with path.open(encoding="utf-8") as handle:
            for raw in handle:
                line = raw.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                key, value = line.split("=", 1)
                values[key] = value.strip().strip("'\"")
    values.update({k: v for k, v in os.environ.items() if k in ("NEXT_PUBLIC_SUPABASE_URL", "NEXT_PUBLIC_SUPABASE_ANON_KEY")})
    return values


def fetch_missing(base_url, anon_key, table, column):
    """Paginiert (statt eines einzelnen limit=1000) — sonst würden Zeilen
    jenseits der ersten 1000 still fehlen, sobald eine Tabelle so groß wird."""
    rows = []
    offset = 0
    headers = {"apikey": anon_key, "Authorization": f"Bearer {anon_key}"}
    while True:
        params = urllib.parse.urlencode({
            "select": f"id,{column}",
            "photo_url": "is.null",
            "order": f"{column}.asc",
            "limit": str(PAGE_SIZE),
            "offset": str(offset),
        })
        request = urllib.request.Request(f"{base_url}/rest/v1/{table}?{params}", headers=headers)
        with urllib.request.urlopen(request, timeout=30) as response:
            page = json.load(response)
        if not page:
            break
        rows.extend(page)
        if len(page) < PAGE_SIZE:
            break
        offset += PAGE_SIZE
    return [row[column] for row in rows]


def set_font(run, size=10.5, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            set_cell_margins(cell)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Seite ")
    set_font(run, 9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_heading(doc, text, count, page_break=False):
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break
    run = paragraph.add_run(text)
    set_font(run, 16, True, BLUE)
    badge = paragraph.add_run(f"  {count}")
    set_font(badge, 10, True, MUTED)


def add_name_table(doc, names):
    rows = (len(names) + 1) // 2
    table = doc.add_table(rows=rows + 1, cols=2)
    table.style = "Table Grid"
    headers = ("Name", "Name")
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_font(run, 9.5, True, WHITE)
    set_repeat_header(table.rows[0])
    for row_idx in range(rows):
        for col_idx in range(2):
            name_idx = row_idx + (rows * col_idx)
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2 == 1:
                shade_cell(cell, "F6F8FA")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if name_idx < len(names):
                set_font(p.add_run(names[name_idx]), 9.5)
    set_table_geometry(table, [4560, 4560])


def add_name_tables(doc, label, names, rows_per_page=18):
    page_capacity = rows_per_page * 2
    chunks = [names[index:index + page_capacity] for index in range(0, len(names), page_capacity)]
    for index, chunk in enumerate(chunks):
        if index:
            continuation = doc.add_paragraph(style="Heading 2")
            continuation.paragraph_format.keep_with_next = True
            continuation.paragraph_format.page_break_before = True
            set_font(continuation.add_run(f"{label} – Fortsetzung"), 13, True, BLUE)
        add_name_table(doc, chunk)


def build_document(base: str, key: str, output_path: Path) -> dict:
    venues = fetch_missing(base, key, "venues", "name")
    persons = fetch_missing(base, key, "persons", "full_name")
    ensembles = fetch_missing(base, key, "ensembles", "name")

    doc = Document()
    # Die mitgelieferte Word-Basistemplate kann Spiegelränder bzw. getrennte
    # Gerade-/Ungerade-Kopfzeilen enthalten. Für diese einseitig gebundene
    # Referenzliste werden beide Eigenschaften explizit entfernt.
    settings = doc.settings._element
    for tag in ("mirrorMargins", "evenAndOddHeaders", "bookFoldPrinting", "bookFoldRevPrinting"):
        node = settings.find(qn(f"w:{tag}"))
        if node is not None:
            settings.remove(node)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.gutter = Inches(0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DARK_BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "KLASSIK MÜNCHEN  ·  BILDLÜCKEN"
    header.paragraph_format.space_after = Pt(0)
    set_font(header.runs[0], 8.5, True, MUTED)
    add_page_field(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(8)
    set_font(kicker.add_run("DATENBANK-REFERENZ"), 9, True, BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("Datensätze ohne Bild"), 27, True, DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(subtitle.add_run("Veranstaltungsorte, Personen und Ensembles"), 13, False, MUTED)

    summary = doc.add_table(rows=1, cols=3)
    summary.style = "Table Grid"
    for idx, (label, count) in enumerate((("Veranstaltungsorte", len(venues)), ("Personen", len(persons)), ("Ensembles", len(ensembles)))):
        cell = summary.rows[0].cells[idx]
        shade_cell(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(f"{count}\n"), 17, True, DARK_BLUE)
        set_font(p.add_run(label), 9, True, MUTED)
    set_table_geometry(summary, [3040, 3040, 3040])

    as_of = doc.add_paragraph()
    as_of.paragraph_format.space_before = Pt(8)
    as_of.paragraph_format.space_after = Pt(8)
    set_font(as_of.add_run(f"Stand: {datetime.now().strftime('%d.%m.%Y')} · Kriterium: photo_url ist leer"), 9, False, MUTED)

    add_heading(doc, "Veranstaltungsorte", len(venues), page_break=True)
    add_name_tables(doc, "Veranstaltungsorte", venues)
    add_heading(doc, "Personen", len(persons), page_break=True)
    add_name_tables(doc, "Personen", persons)
    add_heading(doc, "Ensembles", len(ensembles), page_break=True)
    add_name_tables(doc, "Ensembles", ensembles)

    props = doc.core_properties
    props.title = "Datensätze ohne Bild"
    props.subject = "Veranstaltungsorte, Personen und Ensembles mit leerem photo_url"
    props.author = "Klassik München"
    props.keywords = "Bildlücken, Veranstaltungsorte, Personen, Ensembles"
    doc.save(str(output_path))
    return {"output": str(output_path), "venues": len(venues), "persons": len(persons), "ensembles": len(ensembles)}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--env-file", type=Path, default=None)
    parser.add_argument("--output", type=Path, default=None)
    parser.add_argument("--verbose", action="store_true")
    args = parser.parse_args()

    logging.basicConfig(level=logging.DEBUG if args.verbose else logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

    script_dir = Path(__file__).resolve().parent
    repo_root = script_dir.parents[2]
    env_path = args.env_file or (repo_root / "admin" / ".env.local")
    output_path = args.output or (script_dir / "Bildluecken_Veranstaltungsorte_Personen_Ensembles.docx")

    env = load_env(env_path)
    base = env.get("NEXT_PUBLIC_SUPABASE_URL", "").rstrip("/")
    key = env.get("NEXT_PUBLIC_SUPABASE_ANON_KEY", "")
    if not base or not key:
        logger.error("NEXT_PUBLIC_SUPABASE_URL/NEXT_PUBLIC_SUPABASE_ANON_KEY fehlen (gesucht in %s)", env_path)
        return 2

    result = build_document(base, key, output_path)
    logger.info("Dokument geschrieben: %s", json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
