"""Erzeugt ein Word-Dokument mit allen Venues/Personen/Ensembles ohne
deutsche Biografie/Beschreibung — reine Referenz-Übersicht für die
Redaktion, schreibt nichts in die Datenbank.
"""

import argparse
import json
import logging
import os
import sys
import urllib.parse
import urllib.request
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger("build_missing_bios_doc")

BLUE, DARK, MUTED, LIGHT, WHITE = "2E74B5", "1F4D78", "667085", "E8EEF5", "FFFFFF"
PAGE_SIZE = 1000


def load_env(path: Path) -> dict:
    values = {}
    if path.exists():
        with path.open(encoding="utf-8") as handle:
            for raw in handle:
                line = raw.strip()
                if line and not line.startswith("#") and "=" in line:
                    key, value = line.split("=", 1)
                    values[key] = value.strip().strip("'\"")
    values.update({k: v for k, v in os.environ.items() if k in ("NEXT_PUBLIC_SUPABASE_URL", "NEXT_PUBLIC_SUPABASE_ANON_KEY")})
    return values


def fetch_missing(base, key, table, name_col, bio_col):
    """Paginiert (statt eines einzelnen limit=1000) — sonst würden Zeilen
    jenseits der ersten 1000 still fehlen, sobald eine Tabelle so groß wird."""
    all_rows = []
    offset = 0
    headers = {"apikey": key, "Authorization": f"Bearer {key}"}
    while True:
        params = urllib.parse.urlencode({
            "select": f"id,{name_col},{bio_col}",
            "order": f"{name_col}.asc",
            "limit": str(PAGE_SIZE),
            "offset": str(offset),
        })
        req = urllib.request.Request(f"{base}/rest/v1/{table}?{params}", headers=headers)
        with urllib.request.urlopen(req, timeout=30) as response:
            page = json.load(response)
        if not page:
            break
        all_rows.extend(page)
        if len(page) < PAGE_SIZE:
            break
        offset += PAGE_SIZE
    missing = [
        {"name": row[name_col]}
        for row in all_rows
        if not isinstance(row.get(bio_col), str) or not row[bio_col].strip()
    ]
    return missing, len(all_rows)


def set_font(run, size=10, bold=False, color="000000"):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    # Named renderer-safe override: 9120 DXA plus 120 DXA indent keeps the
    # visible border inside the 9360-DXA content box in Word and LibreOffice.
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    props = table._tbl.tblPr
    width = props.find(qn("w:tblW"))
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = props.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        props.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell_margins(cell)


def repeat_header(row):
    props = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    props.append(node)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Seite ")
    set_font(run, 9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def heading(doc, label, count, page_break=True):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = page_break
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(label), 16, True, BLUE)
    set_font(p.add_run(f"  {count}"), 10, True, MUTED)


def record_table(doc, records):
    row_count = (len(records) + 1) // 2
    table = doc.add_table(rows=row_count + 1, cols=2)
    table.style = "Table Grid"
    for idx, label in enumerate(("Name", "Name")):
        cell = table.rows[0].cells[idx]
        shade(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_font(p.add_run(label), 9.5, True, WHITE)
    repeat_header(table.rows[0])
    for row_idx in range(row_count):
        for pair in range(2):
            rec_idx = row_idx + row_count * pair
            name_cell = table.rows[row_idx + 1].cells[pair]
            name_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2:
                shade(name_cell, "F6F8FA")
            name_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            name_cell.paragraphs[0].paragraph_format.line_spacing = 1.05
            if rec_idx < len(records):
                rec = records[rec_idx]
                set_font(name_cell.paragraphs[0].add_run(rec["name"]), 9.2)
    table_geometry(table, [4560, 4560])


def record_pages(doc, label, records, rows_per_page=18):
    per_page = rows_per_page * 2
    for idx in range(0, len(records), per_page):
        if idx:
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.page_break_before = True
            p.paragraph_format.keep_with_next = True
            set_font(p.add_run(f"{label} – Fortsetzung"), 13, True, BLUE)
        record_table(doc, records[idx:idx + per_page])


def build_document(base: str, key: str, output_path: Path) -> dict:
    venues, venues_total = fetch_missing(base, key, "venues", "name", "description_de")
    persons, persons_total = fetch_missing(base, key, "persons", "full_name", "biography_de")
    ensembles, ensembles_total = fetch_missing(base, key, "ensembles", "name", "description_de")

    doc = Document()
    settings = doc.settings._element
    for tag in ("mirrorMargins", "evenAndOddHeaders", "bookFoldPrinting", "bookFoldRevPrinting"):
        node = settings.find(qn(f"w:{tag}"))
        if node is not None:
            settings.remove(node)
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    section.gutter = Inches(0)

    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(11)
    normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(6), 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DARK),
    ):
        style = doc.styles[name]
        style.font.name, style.font.size, style.font.bold = "Calibri", Pt(size), True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)

    header = section.header.paragraphs[0]
    header.text = "KLASSIK MÜNCHEN  ·  FEHLENDE BIOGRAFIEN"
    header.paragraph_format.space_after = Pt(0)
    set_font(header.runs[0], 8.5, True, MUTED)
    page_number(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before, kicker.paragraph_format.space_after = Pt(8), Pt(8)
    set_font(kicker.add_run("DATENBANK-REFERENZ"), 9, True, BLUE)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("Fehlende Biografien und Beschreibungen"), 25, True, DARK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(subtitle.add_run("Alle Veranstaltungsorte, Personen und Ensembles ohne deutsche Biografie oder Beschreibung"), 12.5, color=MUTED)

    summary = doc.add_table(rows=1, cols=3)
    summary.style = "Table Grid"
    for idx, (label, count, total) in enumerate((
        ("Veranstaltungsorte", len(venues), venues_total),
        ("Personen", len(persons), persons_total),
        ("Ensembles", len(ensembles), ensembles_total),
    )):
        cell = summary.rows[0].cells[idx]
        shade(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment, p.paragraph_format.space_after = WD_ALIGN_PARAGRAPH.CENTER, Pt(0)
        set_font(p.add_run(f"{count}\n"), 17, True, DARK)
        set_font(p.add_run(f"{label}\nvon {total}"), 8.8, True, MUTED)
    table_geometry(summary, [3040, 3040, 3040])

    note = doc.add_paragraph()
    note.paragraph_format.space_before, note.paragraph_format.space_after = Pt(8), Pt(8)
    set_font(note.add_run(f"Stand: {datetime.now().strftime('%d.%m.%Y')} · Aufgeführt sind ausschließlich Einträge mit leerem oder nur aus Leerzeichen bestehendem deutschen Biografie-/Beschreibungsfeld."), 9, color=MUTED)

    heading(doc, "Veranstaltungsorte", len(venues))
    record_pages(doc, "Veranstaltungsorte", venues)
    heading(doc, "Personen", len(persons))
    record_pages(doc, "Personen", persons)
    heading(doc, "Ensembles", len(ensembles))
    record_pages(doc, "Ensembles", ensembles)

    props = doc.core_properties
    props.title = "Fehlende Biografien und Beschreibungen"
    props.subject = "Veranstaltungsorte, Personen und Ensembles mit leerer deutscher Biografie oder Beschreibung"
    props.author = "Klassik München"
    props.keywords = "Biografien, Beschreibungen, Datenpflege, Bildstatus"
    doc.save(str(output_path))
    return {"output": str(output_path), "venues": len(venues), "persons": len(persons), "ensembles": len(ensembles)}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--env-file", type=Path, default=None)
    parser.add_argument("--output", type=Path, default=None)
    parser.add_argument("--verbose", action="store_true")
    args = parser.parse_args()

    logging.basicConfig(level=logging.DEBUG if args.verbose else logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

    script_dir = Path(__file__).resolve().parent
    repo_root = script_dir.parents[2]
    env_path = args.env_file or (repo_root / "admin" / ".env.local")
    output_path = args.output or (script_dir / "Fehlende_Biografien_und_Beschreibungen.docx")

    env = load_env(env_path)
    base = env.get("NEXT_PUBLIC_SUPABASE_URL", "").rstrip("/")
    key = env.get("NEXT_PUBLIC_SUPABASE_ANON_KEY", "")
    if not base or not key:
        logger.error("NEXT_PUBLIC_SUPABASE_URL/NEXT_PUBLIC_SUPABASE_ANON_KEY fehlen (gesucht in %s)", env_path)
        return 2

    result = build_document(base, key, output_path)
    logger.info("Dokument geschrieben: %s", json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
